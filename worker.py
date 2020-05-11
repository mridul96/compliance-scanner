# import module
import multiprocessing
import zipfile
import tarfile
import logging
import pyspark  #pip install pyspark==2.4.4
import boto3    #pip install boto3==1.11.0
import pandas   #pip install pandas==1.0.0
import xlrd     #pip install xlrd==1.2.0
import io
import os
import sys
import ast
import docx     #!pip install python-docx==0.8.10
import uuid
import json
import time
import gzip
import tables   #!pip install tables==3.6.1
import pyxlsb   #!pip install pyxlsb==1.0.6
import shutil
import pickle
import pyarrow  #!pip install pyarrow==0.13.0
import psycopg2 #!pip install psycopg2-binary==2.8.4
import functools

# Extension readable text
_TXT = (".txt",".json",".log",".tsv",".csv",".csv000",".yaml",".yml",".html")
_STR = (".ipynb",".py",".sql",".sh",".bash",".js",".temp",".mf",".css") # text
_GZP = (".gz",".csv.gz",".json.gz",".txt.gz",".log.gz",".tar.gz") # gnu-zip
_ZIP = (".zip",) # compressed-zip
_TAR = (".tar",) # compressed-tar
_ORC = (".ORC",".orc") # compressed-orc
_PKT = (".parquet",) # hadoop-file
_HFI = (".h5",".hdf5") # HDF-file
_EXL = (".xlsx",".xls",".xlsb") # MS-Excel
_DOC = (".docx",) # document
_PKL = (".pkl",) # python-pickle
_EXT = _TXT+_STR+_GZP+_ZIP+_TAR+_ORC+_PKT+_HFI+_EXL+_DOC+_PKL # total extension

def putJSONDisk(File=None, Update=None):
    """
    Function to update JSON file passed as the argument, if the file is
    not a valid JSON it will flush the data and write the new content.

        Arguments - File(or File path), Update

    In case of file not found, it will create file in that location.
    Function will log exception if path is not correct.
    """
    try:
        FLAG = False # return variable
        FILE = File
        DATA = Update
        SIZE = len(DATA)
        JSON = dict()
        # check if DATA is valid dictionary
        if isinstance(DATA, dict) and SIZE:
            # file created if not present
            with open(FILE, "a+") as INFILE:
                try:
                    INFILE.seek(0) # reset cursor
                    JSON = json.load(INFILE)
                    JSON.update(DATA)
                    open(FILE, "w").close() # flush the file
                    json.dump(JSON, INFILE, indent=4) # write updated data
                # error as failed to load JSON
                except ValueError as e:
                    INFILE.truncate() # flush data in memory
                    open(FILE, "w").close()  # flush the file
                    json.dump(DATA, INFILE, indent=4) # write data
                # return true if succeed
                FLAG = True
        else:
            raise Warning("`Dictionary` TypeError or Empty #{}.".format(SIZE))
    # in case of exception
    except Exception as e:
        logging.error("`putJSONDisk`: Failed to save file `{}`.\n{}" \
                      .format(FILE, e))
    # return FLAG
    return(FLAG)

def putFileS3(Bucket=None, File=None, Data=None, Isfile=True, Erase=None,
                Id=None, Dtmp=None):
    """
    Function to save file passed in the argument to S3 bucket. If Flag
    is True file will be fetched from disk else Data will be uploaded.

        Arguments - File, Data, Bucket, Isfile, Erase

    Function will log exception if unable to perform write operation,
     or unable to connect with S3 or Bucket is not a valid name.
    """
    try:
        FLAG   = None # return variable
        BUCKET = Bucket
        FILE   = File
        DATA   = Data
        ISFILE = Isfile
        ERASE  = Erase
        UUID   = Id
        DTMS   = Dtmp
        # define dynamic filename
        FILENAME   = FILE.split("/")[-1]
        PUTOBJECT  = DTMS+"/"+UUID+"_"+FILENAME
        # connect with `s3` using boto3
        S3 = boto3.resource('s3')
        # create folder object in S3
        S3.Object(BUCKET,DTMS+"/").put()
        # upload file in S3 from disk
        if ISFILE is True and os.path.exists(FILE):
            UPLOAD   = open(FILE, "rb")
            RESPONSE = S3.Object(BUCKET, PUTOBJECT).put(Body=UPLOAD)
            FLAG     = (RESPONSE["ResponseMetadata"]["HTTPStatusCode"] == 200)
            UPLOAD.close()
        # upload file in S3 from fileobject
        elif ISFILE is False:
            try:
                UPLOAD = io.BytesIO(json.dumps(DATA, indent=4).encode())
            except Exception as e:
                UPLOAD = io.BytesIO(str.encode(str(DATA)))
                logging.warning("Failed to save as JSON; Now saving as string.")
            RESPONSE = S3.Object(BUCKET, PUTOBJECT).put(Body=UPLOAD)
            FLAG     = (RESPONSE["ResponseMetadata"]["HTTPStatusCode"] == 200)
        else:
            raise FileNotFoundError("File does not exist `{}`.".format(FILE))
        # remove file if upload succeed
        os.remove(FILE) if (ERASE and FLAG and os.path.exists(FILE)) else None
        # log the +ve FLAG
        logging.info("`putFileS3`: `{}` Result saved {} S3: `{}`." \
                     .format(UUID+"_"+FILENAME, FLAG, BUCKET))
    # in case of exception
    except Exception as e:
        logging.error("`putFileS3`: `{}` Result saved {} S3: `{}`.\n{}" \
                     .format(UUID+"_"+FILENAME, FLAG, BUCKET, e))
    # return FLAG
    return(FLAG)

def getJSONDisk(File=None):
    """
    Function to read file in "r" mode as JSON from disk
    and it will raise exception in case not a valid json.

        Arguments - File(or File path)

    Function will log exception if path is not correct.
    """
    try:
        DATA = dict() # return variable
        FILE = File
        # read file as JSON if present
        with open(FILE, "r") as INFILE:
            DATA = json.load(INFILE)
    # in case of exception
    except ValueError as e:
        logging.error("`getJSONDisk`: Invalid JSON `{}`.\n{}" \
                         .format(FILE, e))
    # return dictionary
    return(DATA)

def getFileDisk(File=None):
    """
    Function to read file from disk in "r" mode and return
    the file content as list of elements splited on ",".

        Arguments - File(or File path)

    Function will log exception if path is not correct.
    """
    try:
        DATA = str() # return variable
        FILE = File
        # read file if present
        with open(FILE, "r") as INFILE:
            DATA = INFILE.read()
    # in case of exception
    except Exception as e:
        logging.error("`getFileDisk`: Error 404 PATH `{}`.\n{}" \
                      .format(FILE, e))
    # convert string list into python list
    DATA = [ITEM.strip() for ITEM in DATA.split(",") if len(ITEM.strip())>0]
    logging.info("`{}`: #{}".format(FILE.split("/")[-1], len(DATA)))
    # return data
    return(DATA)

def getConnection(User=None, Password=None, DBname=None, Host=None,
                  Port=None):
    """
    Function to create RedShift cursor on every function call using credential.
    """
    try:
        CURSOR   = None
        REDSHIFT = None
        USER = User
        PASS = Password
        DBN  = DBname
        HOST = Host
        PORT = Port
        # connect with RedShift and get cursor
        REDSHIFT = psycopg2.connect(host=HOST,user=USER,port=PORT,
                                    dbname=DBN,password=PASS)
        CURSOR   = REDSHIFT.cursor()
    # in case of exception
    except Exception as e:
        logging.error("`getRedShift`: Failed to connect with db.\n{}" \
                       .format(e))
    # return db cursor
    return(CURSOR, REDSHIFT)

def getUserPIData(UUID=None, Email=None, Fname=None, Lname=None,
                  Credential=None):
    """
    Function generate the SQL query dynamically using User Information
    (i.e Email, First Name and Last Name) and execute it. It return
    dictionary with column name as key and fetched result as value.

        Arguments -  Cursor, UUID, Email, Fname, Lname,

    Function will log exception no result recived for UUID or failed
    fail to execute query using Cursor.
    """
    try:
        FLAG  = False  # return variable
        VALUE = dict() # return variable
        UUID  = UUID   # return variable
        TEMP  = dict() # return variable
        EMAIL = Email
        FNAME = Fname
        LNAME = Lname
        USER  = Credential[0]
        PASS  = Credential[1]
        DBN   = Credential[2]
        HOST  = Credential[3]
        PORT  = Credential[4]
        CURSOR, REDSHIFT = getConnection(User=USER,Password=PASS,DBname=DBN,
                              Host=HOST,Port=PORT)

        # QUERY01 - TODO
        # `table_name`
        try:
            DATA01 = dict() # return variable
            COL01 = (
            "sample_column_name_1",
            "sample_column_name_2",
            "sample_column_name_3",
            "sample_column_name_4",
            "sample_column_name_5",
            "sample_column_name_6"
            )
            TAB01  = (
            ".table_name"
            )
            CLS01 = "                            \
            lower(check_email) = lower('{}') and \
            lower(check_fname) = lower('{}') and \
            lower(check_lname) = lower('{}')     \
            ".format(EMAIL,FNAME,LNAME)
            # generate query dynamically
            QRY01  = "SELECT "+",".join(COL01)+" FROM "+DBN+TAB01+" WHERE"+CLS01
            # execute query using `CURSOR`
            CURSOR.execute(QRY01)
            RES01 = CURSOR.fetchone()
            # check if result is none
            DATA01 = dict(zip(COL01,RES01)) if not isinstance(RES01,type(None)) else {}
        # in case of exception
        except Exception as e:
            logging.warning("QUERY01 `{}`.".format(e))
            CURSOR, REDSHIFT = getConnection(User=USER,Password=PASS,DBname=DBN,
                                  Host=HOST,Port=PORT)
        # close the connection
        try:
            CURSOR.close()
            REDSHIFT.close()
        # in case of exception
        except Exception as e:
            logging.warning("Not able to close connection.".format(e))

        # create result-list (append) - TODO
        LIST =[DATA01]

        # update temp if not empty
        [TEMP.update(ITEM) if len(ITEM) else None for ITEM in LIST]

        # filer data in dictionary
        DATA = {KEY:VALUE for KEY,VALUE in TEMP.items() if isinstance(VALUE,(str,int,float))
                and len(str(VALUE))>2 and VALUE not in ["unknown","Unknown"]}

        # if temp is not empty
        if not len(DATA) == 0:
            # value for dictionary
            VALUE  = {
            "EMAIL":EMAIL,
            "FNAME":FNAME,
            "LNAME":LNAME,
            "DATA" :DATA
            }
            # return data dictionary and FLAG
            FLAG = True
        else:
            # raise warning for no result
            raise Warning("`No result recived`: UUID `{}`.".format(UUID))
    # in case of exception
    except Exception as e:
        logging.error("`getRedShift`: Failed to execute query.\n{}" \
                      .format(e))
    # return data dictionary and FLAG
    return({UUID:VALUE}, FLAG)


def getRedShift(Path=None, User=None, Password=None, DBname=None,
                Host=None, Port=None):
    """
    Function to connect with RedShift and fetch User Identification
    Information using SQL based on request recived in `UserRequest.json`
    from compliance team.

        Arguments - Path, User, Password, Dbname, Host, Port

    Result is in form of JSON which contains Zip of column name and
    there value in Key,Value pair.
    """
    try:
        DATA = dict() # return variable
        FLAG = False  # return variable
        PATH = Path
        CRED = (User,Password,DBname,Host,Port)
        # check if `UserRequest.json` exists
        if os.path.exists(PATH+"UserRequest.json"):
            # read file in memory
            REQUEST = getJSONDisk(File=PATH+"UserRequest.json")
            # execute `getUserPIData` for each UUID
            for UUID,VALUE in REQUEST.items():
                EMAIL, FNAME, LNAME = VALUE
                USER, FLAG = getUserPIData(UUID=UUID,Email=EMAIL,
                                             Fname=FNAME,Lname=LNAME,
                                               Credential=CRED)
                DATA.update(USER) if FLAG else None
            # set flag
            FLAG = True
        # check if redshift file exist on disk
        elif os.path.exists(PATH+"RedShiftPI.json"):
            # read file in memory
            DATA = getJSONDisk(File=PATH+"RedShiftPI.json")
            # update flag
            FLAG = True if len(DATA) else False
        # in case file does not exist
        else:
            raise FileNotFoundError("`UserRequest.json`/`RedShiftPI.json` \
                                    file not found.")
    # in case of exception
    except Exception as e:
        logging.error("`getRedShift`: Failed to process requests.\n{}" \
                      .format(e))
    # return data and FLAG
    return(DATA, FLAG)

def getS3FileName(Bucket=None):
    """
    Function fetch s3-object for the desired bucket using valid bucket
    name and key should contain a valid extension from "_EXT" list.

        Arguments - Bucket

    Function return dictionary which contain UUID for each s3-object
    and its meta like key, last modified and size uuid is unique value
    generate uuid.uuid4
    """
    try:
        FILELIST = list()  # return variable
        BUCKET   = Bucket
        # connect with `s3` using boto3
        try:
            S3 = boto3.resource('s3')
        # in case of exception
        except Exception as e:
            raise ConnectionError("Failed to connect with S3.")
        # loop over all objects
        for OBJECT in S3.Bucket(BUCKET).objects.all():
            SCLS = OBJECT.storage_class
            OKEY = OBJECT.key
            TIME = OBJECT.last_modified.isoformat() # convert
            SIZE = OBJECT.size
            # if KEY endswith with required extension
            if OKEY.lower().endswith(_EXT) and SIZE>0 and SCLS=="STANDARD":
                KEY   = str(uuid.uuid4())
                VALUE = (BUCKET,OKEY,SIZE,TIME)
                FILELIST.append((KEY,VALUE))
            # do nothing if not in extension
            else:
                None
        # log the bucket final FLAG
        logging.info("Bucket: `{}` #{} files.".format(BUCKET, len(FILELIST)))
    # in case of exception
    except Exception as e:
        logging.error("`getS3FileName`: Failed Bucket `{}`\n{}" \
                      .format(BUCKET, e))
    # return list
    return(FILELIST)

def getS3Directory(Pool=None, Inclusive=[], Exclusive=[], File=None):
    """
    Function fetch all bucket name existing in S3 and compile a final
    list based on Inclusive and Exclusive list passed an argument.

        Arguments - Pool, Inclusive, Exclusive, Path

    Process Pool is used to execute the task in parallel for fetching
    metadata of the buckets in bucket list.
    """
    try:
        TSIZE   = float(0) # return variable
        BUCKETS = list()   # return variable
        FLAG    = False    # return variable
        IBUCKET = Inclusive
        EBUCKET = Exclusive
        FILE    = File
        # connect with `s3` using boto3
        try:
            S3 = boto3.resource('s3')
        # in case of exception
        except Exception as e:
            raise ConnectionError("Failed to connect with S3.")
        # set scanner to Inclusive
        if  "all" in IBUCKET:
            S3 = [OBJECT.name for OBJECT in S3.buckets.all()]
            _BUCKETS = [BUCKET for BUCKET in S3 if BUCKET not in EBUCKET]
        # set scanner to Exclusive
        elif IBUCKET:
            _BUCKETS = IBUCKET
        # set empty for other cases
        else:
            raise ValueError("No request to process.")

        # if `BucketList.json` exists
        if os.path.exists(FILE):
            BUCKETDICT = getJSONDisk(File=FILE)
            BUCKETS = [(KEY,VALUE) for KEY,VALUE in BUCKETDICT.items()]
            TSIZE   = sum([VALUE[2] for KEY, VALUE in BUCKETS])
            logging.warning("`BucketList.json` Found PATH: `{}`".format(FILE))
        # if len of buckets is not null
        elif len(_BUCKETS):
            MAPPER  = Pool.map(getS3FileName, _BUCKETS)
            BUCKETS = [ITEM for SUBLIST in MAPPER for ITEM in SUBLIST]
            TSIZE   = sum([VALUE[2] for KEY, VALUE in BUCKETS])
            BUCKETS = sorted(dict(BUCKETS).items(), key=lambda ITEM:ITEM[1][2])
        #  else do nothing
        else:
            raise ValueError("No request to process.")
        # return True if succeed
        FLAG = True
    # in case of exception
    except Exception as e: # error
        logging.error("`getS3Directory`: Failed to execute.\n{}" \
                      .format(e))
    # return False if failed
    return(BUCKETS, TSIZE, FLAG)

def getS3FileContent(Bucket=None, Key=None, Size=None, Date=None):
    """
    Function fetch data from s3 using bucket name and s3 object key
    based on s3 object size.

        Arguments - Bucket, Key, Size, Date

    If size is smaller than the lower limit then file is stored in
    memory and if size is smaller the upper limit then file is
    downloaded on disk and if size is greater then upper limit file
    is skipped from processing.
    """
    try:
        DATA   = str() # return variable
        FLAG   = None  # return variable
        BUCKET = Bucket
        KEY    = Key
        SIZE   = int(Size)
        LLIMIT = int(1024*1024*1024*12) # x
        ULIMIT = int(1024*1024*1024*48) # 4x
        MATCH  = "."+KEY.split(".")[-1] in _ORC+_HFI
        try:
            S3 = boto3.resource('s3')
        # in case of exception
        except Exception as e:
            raise ConnectionError("Failed to connect with S3.")
        # size smaller then lower limit
        if (SIZE < ULIMIT) and (SIZE < LLIMIT) and not MATCH:
            DATA = S3.Bucket(BUCKET).Object(KEY).get()["Body"].read()
            FLAG = True # set FLAG
        # size greater then lower limit
        elif (SIZE < ULIMIT) and (SIZE > LLIMIT) or MATCH:
            # create temp directory
            UUID = str(uuid.uuid4())
            TEMP = SYSTEM+"TEMP"+SEP+UUID+SEP
            os.makedirs(TEMP)
            # download file
            DATA = TEMP + KEY.split("/")[-1] # create filename
            S3.Bucket(BUCKET).download_file(KEY, DATA)
            FLAG = False # set FLAG
        # exception as file is not under defined limit.
        else:
            DATA = "SIZE #{}".format(ULIMIT)
    # in case of exception
    except Exception as e:
        DATA = "FileFetchFailed"
        logging.critical("`getS3FileContent`: Failed to fetch file `{}`\n{}" \
                         .format(KEY, e))
    # return content and FLAG
    return(DATA, FLAG)

def stringMatchHFI(Search=None, Data=None, Flag=None, Key=None):
    """
    Function to process and identify search values from Hierarchical Data
    Format (HDF) files in memory or from disk based on size of the file.
    """
    try:
        FOUND   = dict() # return variable
        STATUS  = None   # return variable
        CONTENT = Data
        KEY     = Key
        SEARCH  = Search
        # check file on memory
        if Flag is True and len(CONTENT):
            # generate dynamic name
            UUID = str(uuid.uuid4())
            NAME = str(uuid.uuid4())
            TEMP = SYSTEM+"TEMP"+SEP+UUID+SEP
            os.makedirs(TEMP)
            FILE = TEMP+NAME
            # save file to disk
            with open(FILE, "wb") as OUTFILE:
                OUTFILE.write(bytearray(CONTENT))
            # recall function with Flag false
            FOUND, STATUS = stringMatchHFI(Search=SEARCH,Data=FILE,Flag=False,Key=KEY)
        # check file on disk
        elif Flag is False and len(CONTENT):
            HDF = pandas.read_hdf(CONTENT)
            # search keywords in data
            for ID,USER in SEARCH.items():
                PI   = list()
                DATA = USER["DATA"]
                HDF = HDF.apply(lambda x: x.apply(str).str.lower(),axis=1)
                for KEY,VALUE in DATA.items():
                    CHECK = HDF.eq(str(VALUE).lower()).any().any()
                    PI.append(KEY) if CHECK else None
                FOUND.update({ID:PI}) if any(PI) else False
            STATUS = True if any(FOUND) else False
        else:
            STATUS = False
    # in case of exception
    except Exception as e:
        # update the status if any info found till error
        STATUS = True if any(FOUND) else None
        FOUND  = "NonReadableFile" if isinstance(STATUS, type(None)) else FOUND
        logging.error("`stringMatchHFI`: Failed to assert file `{}`.\n{}" \
                      .format(KEY, e))
    # return true if any one is true
    return(FOUND, STATUS)

def stringMatchORC(Search=None, Data=None, Flag=None, Key=None):
    """
    Function to process and identify search values from Highly Compressed
    Apache (ORC) files in memory or from disk based on size of the file.
    """
    try:
        FOUND   = dict() # return variable
        STATUS  = None   # return variable
        CONTENT = Data
        KEY     = Key
        SEARCH  = Search
        # check file on memory
        if Flag is True and len(CONTENT):
            # generate dynamic name
            UUID = str(uuid.uuid4())
            NAME = str(uuid.uuid4())
            TEMP = SYSTEM+"TEMP"+SEP+UUID+SEP
            os.makedirs(TEMP)
            FILE = TEMP+NAME
            # save file to disk
            with open(FILE, "wb") as OUTFILE:
                OUTFILE.write(bytearray(CONTENT))
            # recall function with Flag false
            FOUND, STATUS = stringMatchORC(Search=SEARCH,Data=FILE,Flag=False,Key=KEY)
        # check file on disk
        elif Flag is False and len(CONTENT):
            # define SparkSession
            SPARK = pyspark.sql.SparkSession.builder \
                    .master("local") \
                    .appName("Scanner") \
                    .getOrCreate()
            # convert ORC to pandas
            ORC   = SPARK.read.orc(CONTENT)
            ORCDF = ORC.toPandas()
            # search keywords in data
            for ID,USER in SEARCH.items():
                PI    = list()
                DATA  = USER["DATA"]
                ORCDF = ORCDF.apply(lambda x: x.apply(str).str.lower(),axis=1)
                for KEY,VALUE in DATA.items():
                    CHECK = ORCDF.eq(str(VALUE).lower()).any().any()
                    PI.append(KEY) if CHECK else None
                FOUND.update({ID:PI}) if any(PI) else False
            STATUS = True if any(FOUND) else False
        else:
            STATUS = False
    # in case of exception
    except Exception as e:
        # update the status if any info found till error
        STATUS = True if any(FOUND) else None
        FOUND  = "NonReadableFile" if isinstance(STATUS, type(None)) else FOUND
        logging.error("`stringMatchORC`: Failed to assert file `{}`.\n{}" \
                      .format(KEY, e))
    # return true if any one is true
    return(FOUND, STATUS)

def stringMatchPKT(Search=None, Data=None, Flag=None, Key=None):
    """
    Function to process and identify search values from Apache Parquet
    (.parquet) files in memory or from disk based on size of the file.
    """
    try:
        FOUND   = dict() # return variable
        STATUS  = None   # return variable
        CONTENT = Data
        KEY     = Key
        SEARCH  = Search
        # check file on memory
        if Flag is True and len(CONTENT):
            PARQUET = pandas.read_parquet(io.BytesIO(CONTENT))
            # search keywords in data
            for ID,USER in SEARCH.items():
                PI   = list()
                DATA = USER["DATA"]
                PARQUET = PARQUET.apply(lambda x: x.apply(str).str.lower(),axis=1)
                for KEY,VALUE in DATA.items():
                    CHECK = PARQUET.eq(str(VALUE).lower()).any().any()
                    PI.append(KEY) if CHECK else None
                FOUND.update({ID:PI}) if any(PI) else False
            STATUS = True if any(FOUND) else False
        # check file on disk
        elif Flag is False and len(CONTENT):
            PARQUET = pandas.read_parquet(CONTENT)
            # search keywords in data
            for ID,USER in SEARCH.items():
                PI   = list()
                DATA = USER["DATA"]
                PARQUET = PARQUET.apply(lambda x: x.apply(str).str.lower(),axis=1)
                for KEY,VALUE in DATA.items():
                    CHECK = PARQUET.eq(str(VALUE).lower()).any().any()
                    PI.append(KEY) if CHECK else None
                FOUND.update({ID:PI}) if any(PI) else False
            STATUS = True if any(FOUND) else False
        else:
            STATUS = False
    # in case of exception
    except Exception as e:
        # update the status if any info found till error
        STATUS = True if any(FOUND) else None
        FOUND  = "NonReadableFile" if isinstance(STATUS, type(None)) else FOUND
        logging.error("`stringMatchPKT`: Failed to assert file `{}`.\n{}" \
                      .format(KEY, e))
    # return true if any one is true
    return(FOUND, STATUS)

def stringMatchEXL(Search=None, Data=None, Flag=None, Key=None):
    """
    Function to process and identify search values from MS Excel (xls,
    xlsb,xlsx) files in memory or from disk based on size of the file.
    """
    try:
        FOUND   = dict() # return variable
        STATUS  = None # return variable
        CONTENT = Data
        KEY     = Key
        SEARCH  = Search
        # check file on memory
        if Flag is True and len(CONTENT):
            if KEY.lower().endswith(".xlsb"):
                EXCEL = pandas.ExcelFile(io.BytesIO(CONTENT),engine="pyxlsb")
            else:
                EXCEL = pandas.ExcelFile(io.BytesIO(CONTENT),engine="xlrd")
            # search keywords in data
            for ID,USER in SEARCH.items():
                PI   = list()
                DATA = USER["DATA"]
                for SHEETNAME in EXCEL.sheet_names:
                    DF = EXCEL.parse(SHEETNAME)
                    DF = DF.apply(lambda x: x.apply(str).str.lower(),axis=1)
                    for KEY,VALUE in DATA.items():
                        CHECK = DF.eq(str(VALUE).lower()).any().any()
                        PI.append(KEY) if CHECK else None
                FOUND.update({ID:PI}) if any(PI) else False
            STATUS = True if any(FOUND) else False
        # check file on disk
        elif Flag is False and len(CONTENT):
            if KEY.lower().endswith(".xlsb"):
                EXCEL = pandas.ExcelFile(CONTENT,engine="pyxlsb")
            else:
                EXCEL = pandas.ExcelFile(CONTENT,engine="xlrd")
            # search keywords in data
            for ID,USER in SEARCH.items():
                PI   = list()
                DATA = USER["DATA"]
                for SHEETNAME in EXCEL.sheet_names:
                    DF = EXCEL.parse(SHEETNAME)
                    DF = DF.apply(lambda x: x.apply(str).str.lower(),axis=1)
                    for KEY,VALUE in DATA.items():
                        CHECK = DF.eq(str(VALUE).lower()).any().any()
                        PI.append(KEY) if CHECK else None
                FOUND.update({ID:PI}) if any(PI) else False
            STATUS = True if any(FOUND) else False
        else:
            STATUS = False
    # in case of exception
    except Exception as e:
        # update the status if any info found till error
        STATUS = True if any(FOUND) else None
        FOUND  = "NonReadableFile" if isinstance(STATUS, type(None)) else FOUND
        logging.error("`stringMatchEXL`: Failed to assert file `{}`.\n{}" \
                      .format(KEY, e))
    # return true if any one is true
    return(FOUND, STATUS)

def stringMatchPKL(Search=None, Data=None, Flag=None, Key=None):
    """
    Function to process and identify search values from Python
    pickel (pkl) files in memory or from  disk based on size of the file.
    """
    try:
        FOUND   = dict() # return variable
        STATUS  = None # return variable
        CONTENT = Data
        KEY     = Key
        # check file on memory
        if Flag is True and len(CONTENT):
            DATA = pickle.load(io.BytesIO(CONTENT))
            # call stringMatchSTR
            FOUND, STATUS = stringMatchSTR(Search=Search,Data=DATA,Flag=Flag)
        # check file on disk
        elif Flag is False and len(CONTENT):
            with open(CONTENT, "rb") as INFILE:
                DATA = pickle.load(INFILE)
            # call stringMatchSTR
            FOUND, STATUS = stringMatchSTR(Search=Search,Data=DATA,Flag=Flag)
        else:
            STATUS = False
    # in case of exception
    except Exception as e:
        # update the found based on status
        FOUND  = "NonReadableFile" if isinstance(STATUS, type(None)) else FOUND
        logging.error("`stringMatchPKL`: Failed to assert file `{}`.\n{}" \
                      .format(KEY, e))
    # return true if any one is true
    return(FOUND, STATUS)

def stringMatchDOC(Search=None, Data=None, Flag=None, Key=None):
    """
    Function to process and identify search values from Ms Word
    (doc,docx) files in memory or from disk based on size of the file.
    """
    try:
        FOUND   = dict() # return variable
        STATUS  = None # return variable
        CONTENT = Data
        KEY     = Key
        PARA    = list()
        # check file on memory
        if Flag is True and len(CONTENT):
            DOCUMENT = docx.Document(io.BytesIO(CONTENT))
            # search keywords in data
            for PARAGRAPH in DOCUMENT.paragraphs:
                #join the text into one string
                PARA.append(PARAGRAPH.text)
            # join list into one string
            DATA = " ".join(PARA)
            # call stringMatchSTR
            FOUND, STATUS = stringMatchSTR(Search=Search,Data=DATA,Flag=Flag)
        # check file on disk
        elif Flag is False and len(CONTENT):
            DOCUMENT = docx.Document(CONTENT)
            # search keywords in data
            for PARAGRAPH in DOCUMENT.paragraphs:
                #join the text into one string
                PARA.append(PARAGRAPH.text)
            # join list into one string
            DATA = " ".join(PARA)
            # call stringMatchSTR
            FOUND, STATUS = stringMatchSTR(Search=Search,Data=DATA,Flag=Flag)
        else:
            STATUS = False
    # in case of exception
    except Exception as e:
        # update the found based on status
        FOUND  = "NonReadableFile" if isinstance(STATUS, type(None)) else FOUND
        logging.error("`stringMatchDOC`: Failed to assert file `{}`.\n{}" \
                      .format(KEY, e))
    # return true if any one is true
    return(FOUND, STATUS)

def stringMatchSTR(Search=None, Data=None, Flag=None, Key=None):
    """
    Function to process and identify search values from string readable
    formats and files, in memory or from disk based on size of the file.
    Function read file in chunks of 4 GB to optimze the functionality.
    """
    try:
        FOUND  = dict() # return variable
        STATUS = None # return variable
        try:
            CONTENT = str(Data.decode("utf-8"))
        except:
            CONTENT = str(Data)
        KEY    = Key
        SEARCH = Search
        PI     = list()
        # check file on memory
        if Flag is True and len(CONTENT):
            for ID,USER in SEARCH.items():
                DATA  = USER["DATA"]
                PI = [KEY for KEY,VALUE in DATA.items() if
                            str(VALUE).lower() in CONTENT.lower()]
                FOUND.update({ID:PI}) if any(PI) else False
            STATUS = True if any(FOUND) else False
        # check file on disk
        elif Flag is False and len(CONTENT):
            with open(CONTENT, "r") as INFILE:
                for ID,USER in SEARCH.items():
                    DATA  = USER["DATA"]
                    while True:
                        CONTENT = INFILE.read(int(1024*1024*1024*8))
                        if not CONTENT:
                            break
                        else:
                            None
                        PI.append([KEY for KEY,VALUE in DATA.items() if
                                    str(VALUE).lower() in CONTENT.lower()])
                    FOUND.update({ID:PI}) if any(PI) else False
            STATUS = True if any(FOUND) else False
        else:
            STATUS = False
    # in case of exception
    except Exception as e:
        # update the status if any info found till error
        STATUS = True if any(FOUND) else None
        FOUND  = "NonReadableFile" if isinstance(STATUS, type(None)) else FOUND
        logging.error("`stringMatchSTR`: Failed to assert file `{}`.\n{}" \
                      .format(KEY, e))
    # return status and result
    return(FOUND, STATUS)

def stringMatchGZP(Search=None, Data=None, Flag=None, Key=None):
    """
    Function to process and identify search values from GNU Gzip
    files in memory or from disk based on size of the file.
    """
    try:
        FOUND   = dict() # return variable
        STATUS  = None # return variable
        CONTENT = Data
        KEY     = Key
        # check file on memory
        if Flag is True and len(CONTENT):
            CONTENT = io.BytesIO(CONTENT)
            with gzip.GzipFile(fileobj=CONTENT) as GZIPFILE:
                DATA = GZIPFILE.read()
            # search keywords in data
            if KEY.lower().endswith(".gz"):
                FOUND, STATUS = stringMatchSTR(Key=KEY,Search=Search,
                                               Data=DATA,Flag=True)
            else:
                KEY = KEY[:-3]
                FOUND, STATUS = checkExtension(Key=KEY,Search=Search,
                                               Data=DATA,Flag=True)
        # check file on disk
        elif Flag is False and len(CONTENT):
            with gzip.GzipFile(CONTENT) as GZIPFILE:
                DATA = GZIPFILE.read()
            # search keywords in data
            if KEY.lower().endswith(".gz"):
                FOUND, STATUS = stringMatchSTR(Key=KEY,Search=Search,
                                               Data=DATA,Flag=True)
            else:
                KEY = KEY[:-3]
                FOUND, STATUS = checkExtension(Key=KEY,Search=Search,
                                               Data=DATA,Flag=True)
        else:
            STATUS = False
    # in case of exception
    except Exception as e:
        # update the found based on status
        FOUND  = "NonReadableFile" if isinstance(STATUS, type(None)) else FOUND
        logging.error("`stringMatchGZP`: Failed to assert file `{}`.\n{}" \
                      .format(KEY, e))
    # return true if any one is true
    return(FOUND, STATUS)

def stringMatchTAR(Search=None, Data=None, Flag=None, Key=None):
    """
    Function to process and identify search values from Tar or
    Tarball files in memory or from disk based on size of the file.
    """
    try:
        FOUND   = dict() # return variable
        STATUS  = None # return variable
        CONTENT = Data
        KEY     = Key
        _STATUS = list()
        # check file on memory
        if Flag is True and len(CONTENT):
            TARFILE  = tarfile.open(fileobj=io.BytesIO(CONTENT))
            # execute tar extraction
            FILENAME = TARFILE.getnames()
            # search keywords in data
            for NAME in FILENAME:
                DATA = TARFILE.extractfile(NAME).read()
                TFOUND, TSTATUS = checkExtension(Key=NAME,Search=Search,
                                                 Data=DATA,Flag=True)
                _STATUS.append(TSTATUS) # append the status
                FOUND.update({NAME:TFOUND}) if TSTATUS else None
            STATUS = any(_STATUS)
            TARFILE.close()
        # check file on disk
        elif Flag is False and len(CONTENT):
            TARFILE  = tarfile.open(CONTENT)
            # execute tar extraction
            FILENAME = TARFILE.getnames()
            # search keywords in data
            for NAME in FILENAME:
                DATA = TARFILE.extractfile(NAME).read()
                TFOUND, TSTATUS = checkExtension(Key=NAME,Search=Search,
                                                 Data=DATA,Flag=True)
                _STATUS.append(TSTATUS) # append the status
                FOUND.update({NAME:TFOUND}) if TSTATUS else None
            STATUS = any(_STATUS)
            TARFILE.close()
        else:
            STATUS = False
    # in case of exception
    except Exception as e:
        # update the found based on status
        FOUND  = "NonReadableFile" if isinstance(STATUS, type(None)) else FOUND
        logging.error("`stringMatchTAR`: Failed to assert file `{}`.\n{}" \
                      .format(KEY, e))
    # return true if any one is true
    return(FOUND, STATUS)

def stringMatchZIP(Search=None, Data=None, Flag=None, Key=None):
    """
    Function to process and identify search values from Zip
    files in memory or from disk based on size of the file.
    """
    try:
        FOUND   = dict() # return variable
        STATUS  = None # return variable
        CONTENT = Data
        KEY     = Key
        _STATUS = list()
        # check file on memory
        if Flag is True and len(CONTENT):
            ZIPFILE = zipfile.ZipFile(io.BytesIO(CONTENT)) # convert to ZipFile
            # search keywords in data
            for NAME in ZIPFILE.namelist():
                DATA = ZIPFILE.read(NAME) # read file content
                TFOUND, TSTATUS = checkExtension(Key=NAME,Search=Search,
                                                 Data=DATA,Flag=True)
                _STATUS.append(TSTATUS) # append the status
                FOUND.update({NAME:TFOUND}) if TSTATUS else None
            STATUS = any(_STATUS)
            ZIPFILE.close()
        # check file on disk
        elif Flag is False and len(CONTENT):
            with open(CONTENT, "rb") as INFILE:
                ZIPFILE = zipfile.ZipFile(INFILE)
                # search keywords in data
                for NAME in ZIPFILE.namelist():
                    DATA = ZIPFILE.read(NAME) # read file content
                    TFOUND, TSTATUS = checkExtension(Key=NAME,Search=Search,
                                                     Data=DATA,Flag=True)
                    _STATUS.append(TSTATUS) # append the status
                    FOUND.update({NAME:TFOUND}) if TSTATUS else None
                STATUS = any(_STATUS)
                ZIPFILE.close()
        else:
            STATUS = False
    # in case of exception
    except Exception as e:
        # update the found based on status
        FOUND  = "NonReadableFile" if isinstance(STATUS, type(None)) else FOUND
        logging.error("`stringMatchZIP`: Failed to assert file `{}`.\n{}" \
                      .format(KEY, e))
    # return true if any one is true
    return(FOUND, STATUS)

def checkExtension(Key=None, Search=None, Data=None, Flag=None):
    """
    Function to check extension of the s3 object (i.e. file) and call
    the relevant content search function.xs
    """
    try:
        FOUND  = None # defualt return value
        STATUS = None # defualt return value
        KEY    = Key
        # check the extension
        if KEY.lower().endswith(_TXT+_STR):
            FOUND, STATUS = stringMatchSTR(Search=Search,Data=Data,Flag=Flag,Key=KEY)
        elif KEY.lower().endswith(_GZP):
            FOUND, STATUS = stringMatchGZP(Search=Search,Data=Data,Flag=Flag,Key=KEY)
        elif KEY.lower().endswith(_ORC):
            FOUND, STATUS = stringMatchORC(Search=Search,Data=Data,Flag=Flag,Key=KEY)
        elif KEY.lower().endswith(_HFI):
            FOUND, STATUS = stringMatchHFI(Search=Search,Data=Data,Flag=Flag,Key=KEY)
        elif KEY.lower().endswith(_PKT):
            FOUND, STATUS = stringMatchPKT(Search=Search,Data=Data,Flag=Flag,Key=KEY)
        elif KEY.lower().endswith(_EXL):
            FOUND, STATUS = stringMatchEXL(Search=Search,Data=Data,Flag=Flag,Key=KEY)
        elif KEY.lower().endswith(_TAR):
            FOUND, STATUS = stringMatchTAR(Search=Search,Data=Data,Flag=Flag,Key=KEY)
        elif KEY.lower().endswith(_ZIP):
            FOUND, STATUS = stringMatchZIP(Search=Search,Data=Data,Flag=Flag,Key=KEY)
        elif KEY.lower().endswith(_DOC):
            FOUND, STATUS = stringMatchDOC(Search=Search,Data=Data,Flag=Flag,Key=KEY)
        elif KEY.lower().endswith(_PKL):
            FOUND, STATUS = stringMatchPKL(Search=Search,Data=Data,Flag=Flag,Key=KEY)
        else:
            FOUND, STATUS = "ExtensionNotFound", None
    # in case of exception
    except Exception as e:
        logging.error("`checkExtension`: Failed to assert file `{}`.\n{}" \
                      .format(KEY, e))
    # return the status and found
    return(FOUND, STATUS)

def scannerMain(Skip=None, Search=None, Target=None):
    """
    Function to process the request recived multiprocessing pool
    and call the relavent function to fetch/process data and
    return result back.
    """
    try:
        RESULT = dict()       # return variable
        UUID   = Target[0]    # (UUID : (BUCKET, KEY, SIZE, DATE))
        BUCKET = Target[1][0]
        KEY    = Target[1][1]
        SIZE   = Target[1][2]
        DATE   = Target[1][3]
        SKIP   = Skip
        # log the process and task
        logging.info("--PID:`{}` Bucket:`{}` File:`{}`".format(os.getpid(),BUCKET, KEY))
        # fetch the file-data from s3
        DATA, FLAG = getS3FileContent(Bucket=BUCKET,Key=KEY,Size=SIZE,Date=DATE)
        # check the status of `getS3FileContent`
        if isinstance(FLAG, bool):
            FOUND, STATUS = checkExtension(Key=KEY,Search=Search,Data=DATA,Flag=FLAG)
            # check the status of `checkExtension`
            if isinstance(STATUS, bool):
                RESULT.update({UUID: (BUCKET, KEY, FOUND)}) if STATUS else None
            # if `checkExtension` failed to execute
            elif isinstance(STATUS, type(None)):
                VALUE  = (BUCKET, KEY, SIZE, DATE, FOUND)
                putJSONDisk(File=SKIP, Update=dict({UUID:VALUE}))
                # finally raise the Exception
                raise AssertionError("`Not Assessed`: `{}`".format(VALUE))
            # when no status recived
            else:
                raise ValueError("`scannerMain` received unknown flag `{}`." \
                                 .format(STATUS))
        # if file-download failed
        elif isinstance(FLAG, type(None)):
            VALUE  = (BUCKET, KEY, SIZE, DATE, DATA)
            putJSONDisk(File=SKIP, Update=dict({UUID:VALUE}))
            # finally raise the Exception
            raise AssertionError("`Fetch Failed`: `{}`".format(VALUE))
        # when no status recived
        else:
            raise ValueError("`scannerMain` received unknown flag `{}`." \
                             .format(FLAG))
        # delete the file if it exists and FLAG is True
        os.remove(DATA) if (not FLAG and os.path.exists(DATA)) else None
    # in case of exception
    except Exception as e:
        logging.error("`scannerMain`: Execution Skipped.\n{}".format(e))
    # return result
    return(RESULT)

def getBatch(ITERABLE, COUNT=10):
    """
    Function will generate the batches of iterable using length
    and count value.
    """
    LENGTH = len(ITERABLE)
    # yield the next batch
    for NEXT in range(0, LENGTH, COUNT):
        yield ITERABLE[NEXT:min(NEXT + COUNT, LENGTH)]

if __name__ == '__main__':
    """
    Script `worker.py` designed to auto scale the load over the multiple
    cores of the existing system and perform the following task.

        1. Connect with `s3` to fetch bucket name and s3 objects.
        2. Connect with `redshift` for user PI data.
        3. Connect with `s3` and search PI data in s3 buckets.

    Script save the result to s3 under filename `RedShiftPI.json`,
    `BucketList.json`,`BucketSkip.json` and final result `S3Scanner.json`.

    Note: Script auto clean any file downloaded in memory or on disk.
    """
    # mark script start time
    START = time.time()

    try:
        # define sep and cwd
        CWD, SEP = os.getcwd(), os.sep
        # fetch cmd line arguments
        SYSTEM   =  sys.argv[1] # "./"
    except Exception as e:
        # create path dynamically
        SYSTEM = CWD + SEP + "system" + SEP
    finally:
        TEMP = SYSTEM+"TEMP"+SEP
        # create system directory
        if not os.path.exists(SYSTEM):
            os.makedirs(SYSTEM)

    # unique id for session
    SESS = str(uuid.uuid4())[:8]
    DTMS = time.strftime("%Y%m%d")

    # establish filename
    REQUESTFILE  = SYSTEM+"UserRequest.json"
    REDSHIFTFILE = SYSTEM+"RedShiftPI.json"
    BUCKETFILE   = SYSTEM+"BucketList.json"
    SKIPPEDFILE  = SYSTEM+"BucketSkip.json"
    RESULTFILE   = SYSTEM+"S3Scanner.json"
    LOGFILE      = SYSTEM+"worker.log"

    # logging configuration
    FORMAT = "[%(asctime)s]-%(levelname)s-%(lineno)d-%(process)d---%(message)s"
    logging.basicConfig(format=FORMAT,filename=LOGFILE,level=logging.INFO)

    # log the general information
    CORE = multiprocessing.cpu_count()
    logging.info("CWD Path: `{}`".format(SYSTEM))
    logging.info("CPU Count: #{}".format(CORE))

    # if `CREDENTIAL.json` exists
    if os.path.exists(SYSTEM+"CREDENTIAL.json"):
        CREDENTIAL = getJSONDisk(File=SYSTEM+"CREDENTIAL.json")
        S3SAVE  = CREDENTIAL["S3SAVE"]
        USER    = CREDENTIAL["USER"]
        HOST    = CREDENTIAL["HOST"]
        PORT    = CREDENTIAL["PORT"]
        DBNAME  = CREDENTIAL["DBNAME"]
        PASSWORD= CREDENTIAL["PASSWORD"]
    # if file does not exists
    else:
        logging.critical("`CREDENTIAL.json` Error 404 file not found.")
        sys.exit("`worker.py` execution stopped in between.") # exit system

    # connect and fetch data from redshift
    SEARCH, REDSHIFT = getRedShift(Path=SYSTEM,User=USER,Host=HOST,Port=PORT,
                                   DBname=DBNAME,Password=PASSWORD)

    # check redshift status
    logging.critical("`getRedShift`: Status: `{}`.".format(REDSHIFT))
    None if REDSHIFT else sys.exit("`worker.py` execution stopped in between.")

    # save file to S3
    putFileS3(File=REQUESTFILE,Bucket=S3SAVE,Isfile=True,Erase=True,Id=SESS,Dtmp=DTMS)
    putFileS3(File=REDSHIFTFILE,Bucket=S3SAVE,Isfile=False,Data=SEARCH,Id=SESS,Dtmp=DTMS)

    # fetch data from files
    INCLUSIVE = getFileDisk(SYSTEM+"InclusiveBucket.txt")
    EXCLUSIVE = getFileDisk(SYSTEM+"ExclusiveBucket.txt")

    PROCESS   = CORE - 2 if len(INCLUSIVE) > CORE else len(INCLUSIVE) + 2
    # fetch metadata of S3 buckets
    with multiprocessing.Pool(processes=PROCESS) as pool:
        # execute the `getS3Directory` - multiprocessing.
        BUCKET, SIZE, FINDER = getS3Directory(Pool=pool,File=BUCKETFILE,
                Inclusive=INCLUSIVE,Exclusive=EXCLUSIVE)

    # check finder status
    logging.critical("`Finder`: Status: `{}`.".format(FINDER))
    None if FINDER else sys.exit("`worker.py` execution stopped in between.")

    # save file to S3
    putFileS3(File=BUCKETFILE,Bucket=S3SAVE,Isfile=False,Data=BUCKET,Id=SESS,Dtmp=DTMS)

    # mark script end time
    MARK = time.time()

    # log the general information
    logging.info("Total File Size  : #{0:.2f} GB.".format(SIZE/(1024*1024*1024)))
    logging.info("Total S3 Object  : #{}.".format(len(BUCKET)))
    logging.info("`scheduler` execution completed in {0:.2f} min." \
                 .format((MARK-START)/60))

    # scanner configuration
    BATCHNUMBER = 1
    PROCESS     = CORE - 2 if CORE < 16 else CORE - 4
    BATCHSIZE   = PROCESS*10
    logging.info("Batch Size #{}.".format(BATCHSIZE))

    # execute the scanner
    for NEXTBATCH in getBatch(BUCKET[:], BATCHSIZE):
        T1 = time.time() # batch start time
        TSIZE = sum([VALUE[2] for KEY, VALUE in NEXTBATCH])
        # log batch size in GB
        logging.info("Batch {0:} started processing {1:.2f} GB data." \
                     .format(BATCHNUMBER, TSIZE/(1024*1024*1024)))
        # execute the batch - multiprocessing.
        with multiprocessing.Pool(processes=PROCESS) as pool:
            FUNC   = functools.partial(scannerMain, SKIPPEDFILE, SEARCH)
            OUTPUT = pool.map(FUNC, NEXTBATCH)
        # save the JSON result
        DATA = dict((KEY, ITEM[KEY]) for ITEM in OUTPUT for KEY in ITEM)
        # define batch filename
        BATCHFILE = "Batch-{}".format(BATCHNUMBER)+".json"
        # save data if received
        if len(DATA):
            # save file to S3
            putFileS3(File=BATCHFILE,Bucket=S3SAVE,Isfile=False,Data=DATA,Id=SESS,Dtmp=DTMS)
            # save to disk
            STATUS = putJSONDisk(File=RESULTFILE, Update=DATA)
        else:
            STATUS = None
        T2 = time.time() # batch end time
        # log the batch execution time
        logging.info("Batch {0:} completed `{1:}` in {2:.2f} min." \
                     .format(BATCHNUMBER,STATUS,(T2-T1)/60))
        # increment batch number
        BATCHNUMBER = BATCHNUMBER + 1

    # save file to S3
    putFileS3(File=SKIPPEDFILE,Bucket=S3SAVE,Isfile=True,Erase=True,Id=SESS,Dtmp=DTMS)
    putFileS3(File=RESULTFILE,Bucket=S3SAVE,Isfile=True,Erase=False,Id=SESS,Dtmp=DTMS)

    # remove temporary folder
    shutil.rmtree(path=TEMP, ignore_errors=True)

    # mark script end time
    STOP = time.time()

    # log the general information
    logging.info("`worker.py` execution completed in {0:.2f} min." \
                 .format((STOP-MARK)/60))

    # save file to S3
    putFileS3(File=LOGFILE,Bucket=S3SAVE,Isfile=True,Erase=False,Id=SESS,Dtmp=DTMS)

    sys.exit("`worker.py` execution completed.")
